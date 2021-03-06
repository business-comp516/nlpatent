# -*- coding: utf-8 -*-

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from docx.shared import Pt
from gen_const import BILLING_PATH, HST_ACCOUNT_NUMBER
from gen_const import calculate_due_date

##### HELPERS #####

def add_styles(styles, txtFont, txtSize):
    # this changes the normal text font & size
    style = styles['Normal']
    font = style.font
    font.name = txtFont
    font.size = Pt(txtSize)
    para_format = style.paragraph_format
    para_format.space_after = Pt(0)
    # OTHER OPTIONS: font.bold, font.italic, font.underline = True, True, True

    # this changes the style of 'Quote'
    style = styles['Quote']
    paragraph_format = style.paragraph_format
    paragraph_format.left_indent = Inches(0.25)
    #paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # or LEFT/RIGHT/JUSTIFY

    # this defines a new style for the Title (bug: can't change style when doing .add_heading)
    new_title_style = styles.add_style('NewTitle', WD_STYLE_TYPE.PARAGRAPH)
    new_title_style.base_style = styles['Title']
    font = new_title_style.font
    font.name = txtFont
    font.size = Pt(26)

    # this defines a new style for Heading 1 (bug: can't change style when doing .add_heading)
    new_heading_style = styles.add_style('NewHeading1', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']
    font = new_heading_style.font
    font.name = txtFont
    font.size = Pt(txtSize+3)
    para_format = new_heading_style.paragraph_format
    para_format.space_before = Pt(0)
    para_format.space_after = Pt(2)

    # this defines a new style for Heading 2 (bug: can't change style when doing .add_heading)
    new_heading_style = styles.add_style('NewHeading2', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 2']
    font = new_heading_style.font
    font.name = txtFont
    font.size = Pt(txtSize+2)
    para_format = new_heading_style.paragraph_format
    para_format.space_before = Pt(0)
    para_format.space_after = Pt(2)

    # this defines a new style for Heading 3 (bug: can't change style when doing .add_heading)
    new_heading_style = styles.add_style('NewHeading3', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 3']
    font = new_heading_style.font
    font.name = txtFont
    font.size = Pt(txtSize+1)
    para_format = new_heading_style.paragraph_format
    para_format.space_after = Pt(2)

    # this defines a new style for Heading 4 (bug: can't change style when doing .add_heading)
    new_heading_style = styles.add_style('NewHeading4', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 4']
    font = new_heading_style.font
    font.name = txtFont
    font.size = Pt(txtSize)
    font.italic = True
    para_format = new_heading_style.paragraph_format
    para_format.space_after = Pt(2)

    # this defines a new style for the Input Text
    new_style = styles.add_style('InputText', WD_STYLE_TYPE.PARAGRAPH)
    new_style.base_style = styles['Normal']
    font = new_style.font
    font.size = Pt(9)

##### SECONDARY #####

def run(c_id, c, m, d, s, u, s_cost, u_cost, t, amt_usage, amt_tax, amt_total, currency, address, due_date, activation_fee, activation_fee_tax, bank_account):
    # CREATE A DOCUMENT & SET THE STYLES
    document = Document()
    styles = document.styles
    add_styles(styles, 'Arial', 11)

    # ADD TITLE
    document.add_paragraph('NLPatent Invoice', style='NewTitle')

    # ADD LEGALICITY INFORMATION
    p = document.add_paragraph()
    p.add_run(u'Legalicity Ltd.\n26 Blyth St.,\nRichmond Hill, ON, L4E 2Y1').bold = True

    # ADD DIRECT DEPOSIT INFORMATION
    p = document.add_paragraph()
    p = document.add_paragraph()
    p.add_run(u'Pay by direct deposit to ')
    p.add_run(u'{}'.format(bank_account[0])).bold = True
    p.add_run(u' (transit #) - ')
    p.add_run(u'{}'.format(bank_account[1])).bold = True
    p.add_run(u' (fin. inst.) - ')
    p.add_run(u'{}'.format(bank_account[2])).bold = True
    p.add_run(u' (account #)')

    # ADD CLIENT ORGANIZATION & ADDRESS
    p = document.add_paragraph('Client', style='NewHeading3')
    p = document.add_paragraph()
    p.add_run(u'{}'.format(c))
    p = document.add_paragraph()
    p.add_run(u'{}'.format(address))

    # ADD INVOICE DATE
    p = document.add_paragraph('Invoice Date', style='NewHeading3')
    p = document.add_paragraph()
    p.add_run(u'{}'.format(d))

    # ADD DUE DATE
    p = document.add_paragraph('Due Date', style='NewHeading3')
    p = document.add_paragraph()
    p.add_run(u'{}'.format(calculate_due_date(d, due_date)))

    # ADD INVOICE NUMBER
    p = document.add_paragraph('Invoice Number', style='NewHeading3')
    p = document.add_paragraph()
    p.add_run(u'{}-{}'.format(c_id.upper(), m))

    # ADD HST ACCOUNT
    p = document.add_paragraph('HST Account Number', style='NewHeading3')
    p = document.add_paragraph()
    p.add_run(u'{}'.format(HST_ACCOUNT_NUMBER))

    # ADD SUMMARY
    p = document.add_paragraph('Summary', style='NewHeading3')
    p = document.add_paragraph()
    # add the activation fee, if any
    if activation_fee:
        p.add_run(u'NLPatent activation fee\t\t\t\t')
        p.add_run(u'{}'.format('${:,.2f}'.format(activation_fee))).bold = True
        # add tax if there is any
        if t[0]:
            p = document.add_paragraph()
            p.add_run(u'{} ({}%)\t\t\t\t\t\t'.format(t[0], t[1]))
            p.add_run(u'{}'.format('${:,.2f}'.format(activation_fee_tax))).bold = True
        p = document.add_paragraph()
        p = document.add_paragraph()
    p.add_run(u'NLPatent usage fee for {}\t\t\t'.format(m))
    p.add_run(u'{}'.format('${:,.2f}'.format(amt_usage))).bold = True
    p = document.add_paragraph()
    p.add_run(u'(Searches: {} x ${})'.format(s, s_cost))
    p = document.add_paragraph()
    p.add_run(u'(Updates: {} x ${})'.format(u, u_cost))
    # add tax if there is any
    if t[0]:
        p = document.add_paragraph()
        p.add_run(u'{} ({}%)\t\t\t\t\t\t'.format(t[0], t[1]))
        p.add_run(u'{}'.format('${:,.2f}'.format(amt_tax))).bold = True
    p = document.add_paragraph()
    p = document.add_paragraph()
    p.add_run(u'Amount Due\t\t\t\t\t\t')
    p.add_run(u'{} {}'.format('${:,.2f}'.format(amt_total), currency)).bold = True

    # SAVE DOCUMENT AS WORD FILE
    f_name = '{}_invoice'.format(c_id)
    doc_name_word = '{}/{}/{}.docx'.format(BILLING_PATH, m, f_name)
    document.save(doc_name_word)

